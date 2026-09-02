"""Generate Home Dashboard reference Word document."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Home-Dashboard-Reference.docx"


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, "Home Dashboard — Reference Guide")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Generated {date.today().isoformat()}")
    run.italic = True
    run.font.size = Pt(11)
    doc.add_paragraph()

    add_para(
        doc,
        "This document describes the Home Dashboard web application deployed to Home Assistant "
        "(/config/www/home-dashboard/). It covers every home-page widget, drill-down page, "
        "Home Assistant automation, integration, configuration file, and setup step required "
        "to operate the system.",
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "1. Overview")
    add_bullets(
        doc,
        [
            "Full control URL: http://homeassistant.local:8123/local/home-dashboard/index.html",
            "Read-only URL: http://homeassistant.local:8123/local/home-dashboard/view.html",
            "Deploy command: npm run deploy (builds and copies to HA via Samba)",
            "Primary data source: Home Assistant REST API (5-second poll; Crestron lights 2-second poll)",
            "Secondary data: JSON cache files on HA, Zynect cloud API, Homebridge schedule",
        ],
    )

    add_heading(doc, "2. Initial Deployment & Setup", level=2)
    add_heading(doc, "2.1 One-time deploy setup", level=3)
    add_numbered(
        doc,
        [
            "Enable Samba on Home Assistant (Settings → System → Storage).",
            "Copy deploy.env.example to deploy.env (default share: \\\\homeassistant.local\\config).",
            "Copy public\\ha-config.example.json to ha-config.json at the project root.",
            "Create a Home Assistant long-lived access token (Profile → Long-Lived Access Tokens) and paste it into ha-config.json. Leave baseUrl empty when the dashboard runs on the same HA host.",
            "Run npm run deploy from the project root.",
            "Restart Home Assistant after the first deploy that includes custom components.",
            "Open the dashboard URL and hard-refresh if an old cached version appears.",
        ],
    )

    add_heading(doc, "2.2 Home Assistant configuration (choose one path)", level=3)
    add_para(doc, "Option A — Packages (recommended if no conflicts):", bold=True)
    add_para(
        doc,
        "Add under homeassistant: in configuration.yaml:\n"
        "  packages: !include_dir_named packages\n"
        "Then restart or reload YAML. Deploy copies packages/home-dashboard.yaml automatically.",
    )
    add_para(doc, "Option B — Manual snippets:", bold=True)
    add_para(
        doc,
        "Follow homeassistant/snippets/README.md: merge shell commands, scripts, input helpers, "
        "and automations into configuration.yaml, scripts.yaml, and automations.yaml. "
        "This project’s live HA instance uses Option B for outside-light and shed-power automations.",
    )

    add_heading(doc, "2.3 Optional cache-sync automation", level=3)
    add_para(
        doc,
        "For remote dashboards and accurate shade positions without LAN access to Homebridge, "
        "add the automation documented in homeassistant/dashboard_sync/README.md. "
        "It runs sync_caches.py every 15 minutes and on HA start, writing shade-schedule-today.json "
        "and shades-cache.json.",
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "3. Dashboard Modes")
    add_table(
        doc,
        ["Mode", "URL", "Behavior"],
        [
            [
                "Full control",
                "index.html",
                "All toggles, sliders, mode buttons, Settings page, and mapping tools enabled.",
            ],
            [
                "Read-only",
                "view.html",
                "Live data only. Controls disabled. Settings route redirects home. "
                "Shows a “View only” badge in the header.",
            ],
        ],
    )
    add_para(
        doc,
        "Both modes load assets via version.json for cache busting. "
        "Read-only detection is in src/dashboardMode.ts.",
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "4. Home Page Widgets")

    widgets = [
        (
            "Weather",
            "WeatherWidget",
            "Current temperature, conditions, wind, and 5-day forecast.",
            "Home Assistant weather.* entity (prefers weather.home, weather.forecast_home, weather.buienradar).",
            "HA weather integration (e.g. Met.no, OpenWeatherMap, or HA default). No dashboard mapping required.",
        ),
        (
            "Shades overview",
            "ShadesOverviewWidget",
            "Summary of 33 shades with per-floor compass showing closed ratio by direction.",
            "HA cover.* entities via shade-map.json; optional shades-cache.json override.",
            "Map each shade to a cover entity in Settings. Link to /shades for full controls.",
        ),
        (
            "Shed Solar (combined solar widget)",
            "ShedSolarWidget",
            "Three panes: Shed (Enphase PowerPack + grid toggle), PV Array (AlsoEnergy), "
            "and Solar Thermal (Zynect summary + sun arc).",
            "energy-map.json sensors; pv-cache.json; shed-cache.json; zynect-config.json; switch.shed_power.",
            "AlsoEnergy + Enphase PowerPack custom components; Zynect credentials; energy sensor mapping in Settings.",
        ),
        (
            "Pool",
            "PoolWidget",
            "Temperature, spa heater status (Heating/Standby), pump RPM, water level, pool lights toggle.",
            "pool-map.json → Pentair climate/sensor/light entities + YoLink depth sensor.",
            "Pentair ScreenLogic integration; YoLink depth sensor; map entities in Settings.",
        ),
        (
            "Pond",
            "PondWidget",
            "Pond temperature, level %, and adjusted water level.",
            "pond-map.json (Tuya sensors) + Zynect API sensor named “Pond”.",
            "Tuya pond sensors mapped in Settings; Zynect token in zynect-config.json.",
        ),
        (
            "HVAC",
            "HvacWidget",
            "Count of thermostats actively heating and mini-split A/C units cooling.",
            "All HA climate.* entities classified as Nest thermostats or mini-splits.",
            "Google Nest integration for heat; Mitsubishi Kumo (ha_kumo_ws) for cooling.",
        ),
        (
            "Outside",
            "OutsideWidget",
            "Mode buttons (None / Normal / Guest) and Driveway/Pond transformer controls with dimmers.",
            "Hardcoded switch/light entities in src/ha/outside.ts + input_select.outside_lights_mode.",
            "TP-Link or compatible switches; outside-light scripts and automations on HA.",
        ),
        (
            "Lights",
            "LightsWidget",
            "On/total counts per floor; Crestron scene buttons (Guest, Games, Sleep, All off).",
            "Crestron Home entities (platform crestron_home or homekit_controller) + lights-map.json room assignments.",
            "Crestron Home HA integration; optional room setup on Lights page.",
        ),
    ]

    for name, component, displays, source, setup in widgets:
        add_heading(doc, f"4.{widgets.index((name, component, displays, source, setup)) + 1} {name}", level=2)
        add_table(
            doc,
            ["Property", "Detail"],
            [
                ["Component", component],
                ["Displays", displays],
                ["Data source", source],
                ["Setup / integration", setup],
            ],
        )

    add_heading(doc, "4.9 Sun arc (inside Shed Solar widget)", level=2)
    add_para(
        doc,
        "SunArcGraphic shows sunrise-to-sunset progress using site latitude/longitude from "
        "zynect-config.json and HA sun.sun state.",
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "5. Drill-Down Pages")
    add_table(
        doc,
        ["Route", "Page", "Purpose", "Key integrations"],
        [
            ["/shades", "ShadesPage", "Per-shade open/close and position; floor Open/Close All", "HA cover entities; Homebridge schedules"],
            ["/lights", "LightsPage", "Crestron lights by floor and room; brightness; room setup", "Crestron Home; lights-map.json"],
            ["/lights/:groupId", "LightsPage", "Single-floor lights (upstairs, main, basement, outside)", "Same as above"],
            ["/solar-thermal", "SolarThermalPage", "Zynect gauges, heating mode, history charts", "Zynect Thermote API"],
            ["/hvac", "HvacPage", "Nest thermostat tiles by floor; setpoint and mode control", "Google Nest climate entities"],
            ["/ac", "AcPage", "Mitsubishi mini-split tiles by floor; cooling control", "ha_kumo_ws climate entities"],
            ["/settings", "SettingsPage", "Connection, mappings, thresholds (full mode only)", "HA token; all map JSON files"],
        ],
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "6. Home Assistant Automations")
    add_para(
        doc,
        "Automations live in homeassistant/packages/home-dashboard.yaml (package path) "
        "and corresponding snippets. On restart, outside-light automations must NOT use "
        "homeassistant.start triggers — only explicit mode changes and schedule triggers should run scripts.",
    )

    automations = [
        (
            "shed_power_initialize_thresholds",
            "Initialize Shed Power thresholds",
            "Home Assistant start (once)",
            "If input_boolean.shed_power_thresholds_initialized is off, set on-threshold to 20%, off-threshold to 80%, then mark initialized.",
        ),
        (
            "shed_power_on_low_battery",
            "Shed Power turn on when SOC falls below threshold",
            "State change on sensor.enphase_powerpack_5904582_battery_soc",
            "When SOC crosses downward through input_number.shed_power_on_soc_threshold, turn on switch.shed_power.",
        ),
        (
            "shed_power_off_high_battery",
            "Shed Power turn off when SOC rises above threshold",
            "State change on sensor.enphase_powerpack_5904582_battery_soc",
            "When SOC crosses upward through input_number.shed_power_off_soc_threshold, turn off switch.shed_power.",
        ),
        (
            "dashboard_persist_map_offset",
            "Dashboard persist pool/pond map offset",
            "Event dashboard_update_map_offset",
            "Runs shell_command.dashboard_update_map_offset to write depthOffset into pool-map.json or pond-map.json.",
        ),
        (
            "outside_lights_none_mode",
            "Outside lights None mode",
            "input_select.outside_lights_mode changes to None (user action only)",
            "Runs script.outside_lights_all_off.",
        ),
        (
            "outside_lights_normal_mode",
            "Outside lights Normal mode",
            "Mode → Normal; sunset; sunrise; daily 22:30",
            "At sunset or when user selects Normal before 22:30: normal_on (sign, upper driveway, courtyard, garden). "
            "At 22:30: normal_late_off. At sunrise: sign_off. Does not run on HA startup.",
        ),
        (
            "outside_lights_guest_mode",
            "Outside lights Guest mode",
            "input_select.outside_lights_mode changes to Guest (user action only)",
            "Runs script.outside_lights_guest_on (all outside lights on).",
        ),
        (
            "dashboard_cache_sync_periodic (optional)",
            "Sync dashboard caches",
            "Every 15 minutes; HA start",
            "Runs sync_caches.py for Homebridge schedule and shade position caches.",
        ),
    ]
    add_table(
        doc,
        ["Automation ID", "Purpose", "Trigger", "Action summary"],
        [[a[0], a[1], a[2], a[3]] for a in automations],
    )

    add_heading(doc, "6.1 Outside light scripts", level=2)
    add_table(
        doc,
        ["Script", "What it does"],
        [
            ["outside_lights_all_off", "Turn off gate, west side, pond switches and driveway light."],
            [
                "outside_lights_normal_on",
                "All off first, then on: sign (gate_switch_2), upper driveway, courtyard, garden.",
            ],
            [
                "outside_lights_normal_late_off",
                "Turn off upper driveway, courtyard, garden, pond switches, and driveway light.",
            ],
            ["outside_lights_sign_off", "Turn off sign only (switch.gate_switch_2)."],
            ["outside_lights_guest_on", "Turn on all outside switches and driveway light."],
        ],
    )

    add_heading(doc, "6.2 Outside light entity map", level=2)
    add_table(
        doc,
        ["Control label", "Entity ID(s)", "Notes"],
        [
            ["Sign", "switch.gate_switch_2", "Switch"],
            ["Lower Driveway", "light.driveway_lights_light_1", "Dimmable light"],
            ["Upper Driveway", "switch.gate_switch_1, switch.west_side_switch_1", "Two switches"],
            ["Courtyard lights", "switch.west_side_switch_2", "Switch"],
            ["Pond", "switch.pond_switch_2", "Switch"],
            ["Garden", "switch.pond_switch_1", "Switch"],
        ],
    )
    add_para(
        doc,
        "Mode helper: input_select.outside_lights_mode — options None, Normal, Guest. "
        "Persisted by Home Assistant; dashboard reads and writes via select_option service.",
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "7. Integrations & Custom Components")
    add_table(
        doc,
        ["Integration", "Type", "Used for", "Install / notes"],
        [
            [
                "Home Assistant core",
                "Local API",
                "All live control and state",
                "Long-lived token in ha-config.json",
            ],
            [
                "Lutron shades (covers)",
                "HA cover.*",
                "33 window shades",
                "Mapped in shade-map.json",
            ],
            [
                "Homebridge schedule plugin",
                "External HTTP :8787",
                "Shade open/close schedules",
                "Fetched at build + optional sync_caches.py",
            ],
            [
                "Crestron Home",
                "HA crestron_home",
                "Interior/exterior lights and scenes",
                "Entity registry platform filter",
            ],
            [
                "AlsoEnergy PowerTrack",
                "Custom component (cloud)",
                "PV array production and energy",
                "custom_components/alsoenergy; poll 10 min daylight; writes pv-cache.json",
            ],
            [
                "Enphase PowerPack (Cloud)",
                "Custom component (cloud)",
                "Shed solar, battery SOC, load, grid",
                "custom_components/enphase_powerpack; poll 2 min; writes shed-cache.json",
            ],
            [
                "TP-Link Kasa",
                "HA switch",
                "Shed grid power outlet",
                "switch.shed_power",
            ],
            [
                "Pentair ScreenLogic",
                "HA climate/sensor/light",
                "Pool temp, pump, spa heat, SAm lights",
                "Mapped in pool-map.json",
            ],
            [
                "YoLink",
                "HA sensor",
                "Pool water depth",
                "sensor.water_depth_sensor_distance",
            ],
            [
                "Tuya",
                "HA sensor",
                "Pond level and depth",
                "sensor.pond_level_liquid_level, sensor.pond_level_depth",
            ],
            [
                "Google Nest",
                "HA climate",
                "Room heating thermostats",
                "Excluded from mini-split detection in hvac.ts",
            ],
            [
                "Mitsubishi Kumo (ha_kumo_ws)",
                "HA climate",
                "Mini-split A/C cooling",
                "Detected by serial, swing_modes, dry mode",
            ],
            [
                "Zynect / Thermote",
                "Direct HTTPS API",
                "Solar thermal temps + pond temperature",
                "Bearer token in zynect-config.json; not HA entities",
            ],
            [
                "Weather",
                "HA weather.*",
                "Home page weather widget",
                "Any HA weather integration",
            ],
        ],
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "8. Configuration Files")
    add_para(
        doc,
        "All runtime JSON files live on the HA host at /config/www/home-dashboard/. "
        "Deploy preserves existing user config files and never overwrites tokens or maps.",
    )
    add_table(
        doc,
        ["File", "Purpose", "Written by"],
        [
            ["ha-config.json", "HA token and optional baseUrl", "User / deploy seed"],
            ["shade-map.json", "Shade ID → cover entity (33 shades)", "Settings export"],
            ["energy-map.json", "Solar sensor entity mappings", "Settings export"],
            ["pool-map.json", "Pool sensors + depthOffset (inches)", "Settings / HA script"],
            ["pond-map.json", "Pond sensors + depthOffset", "Settings / HA script"],
            ["lights-map.json", "Crestron light → room assignments", "HA script dashboard_set_crestron_light_room"],
            ["zynect-config.json", "Zynect auth + site coordinates + refresh interval", "Settings export"],
            ["pv-cache.json", "AlsoEnergy PV snapshot", "alsoenergy custom component"],
            ["shed-cache.json", "Enphase PowerPack snapshot", "enphase_powerpack custom component"],
            ["shades-cache.json", "Cover positions for remote dashboards", "sync_caches.py"],
            ["shade-schedule-today.json", "Homebridge schedule for today", "Build / sync_caches.py"],
            ["homebridge-schedule.json", "Schedule name → shade ID aliases", "Deploy from public/"],
            ["version.json", "Cache-bust loader metadata", "Vite build"],
        ],
    )

    add_heading(doc, "8.1 Example energy-map.json entities", level=2)
    add_bullets(
        doc,
        [
            "pvOnlyProduction: sensor.pv_66885_production_power",
            "pvOnlyTodayEnergy / Month / Lifetime: sensor.pv_66885_energy_produced_*",
            "powerpackProduction: sensor.enphase_powerpack_5904582_pv_production_power",
            "powerpackBatterySoc: sensor.enphase_powerpack_5904582_battery_soc",
            "powerpackLoad / BatteryPower / Grid: sensor.enphase_powerpack_5904582_*",
        ],
    )

    add_heading(doc, "8.2 HA input helpers", level=2)
    add_table(
        doc,
        ["Helper", "Entity ID", "Default", "Purpose"],
        [
            ["input_number", "input_number.shed_power_on_soc_threshold", "20%", "Turn on shed grid below this SOC"],
            ["input_number", "input_number.shed_power_off_soc_threshold", "80%", "Turn off shed grid above this SOC"],
            ["input_boolean", "input_boolean.shed_power_thresholds_initialized", "off", "One-time threshold init flag"],
            ["input_select", "input_select.outside_lights_mode", "None", "Outside lighting mode"],
        ],
    )

    add_heading(doc, "8.3 HA shell commands & scripts", level=2)
    add_table(
        doc,
        ["Name", "Purpose"],
        [
            ["shell_command.dashboard_update_map_offset", "Write pool/pond depth offset to JSON"],
            ["shell_command.dashboard_update_crestron_light_room", "Write Crestron room assignment to lights-map.json"],
            ["shell_command.sync_dashboard_caches", "Run sync_caches.py (optional automation)"],
            ["script.dashboard_set_map_offset", "Wrapper for map offset shell command"],
            ["script.dashboard_set_crestron_light_room", "Wrapper for Crestron room shell command"],
        ],
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "9. Settings Page")
    add_para(doc, "Available at /settings in full control mode only.")
    add_table(
        doc,
        ["Section", "Configures"],
        [
            ["Connection", "HA token, base URL, connect/disconnect, manual refresh"],
            ["Share mappings", "Export/import ha-config and all map JSON files"],
            ["Shed Power thresholds", "input_number shed on/off SOC thresholds (synced to HA)"],
            ["Zynect — Solar thermal", "API auth header, site lat/long, refresh interval"],
            ["Pool / Pond depth offset", "depthOffset in pool-map.json / pond-map.json via HA event"],
            ["Solar sensor mapping", "PV Array (AlsoEnergy) and Shed Solar (Enphase) entity dropdowns"],
            ["Shade schedule debug", "Homebridge cache stats, HA automation load diagnostics"],
            ["Entity map — Shade → cover", "Per-shade cover entity assignment; auto-match by name"],
        ],
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "10. External Services")
    add_table(
        doc,
        ["Service", "Endpoint", "Role"],
        [
            ["Home Assistant", "Same origin or baseUrl", "Primary data and control"],
            ["Homebridge schedule", "http://homebridge.local:8787/schedule.json", "Shade schedules"],
            ["Samba share", "\\\\homeassistant.local\\config", "Windows deploy path"],
            ["Zynect API", "https://zynect.com/api/v2/", "Solar thermal + pond temperature"],
            ["AlsoEnergy API", "https://api.alsoenergy.com", "PV production via custom component"],
            ["Enphase Enlighten", "https://enlighten.enphaseenergy.com", "Shed PowerPack via custom component"],
            ["Nabu Casa (optional)", "Remote HA URL", "Same dashboard at /local/home-dashboard/"],
        ],
    )

    # ------------------------------------------------------------------ #
    add_heading(doc, "11. Feature Setup Checklist")
    checklist = [
        ("Shades", "Map 33 shades to cover.* entities; export shade-map.json; optional cache-sync automation."),
        ("Solar PV (array)", "Install AlsoEnergy integration; map sensors in Settings; verify pv-cache.json updates."),
        ("Shed solar", "Install Enphase PowerPack (site 5904582); map sensors; configure SOC thresholds; verify shed-cache.json."),
        ("Shed grid toggle", "Ensure switch.shed_power exists (TP-Link Kasa); automations enabled."),
        ("Solar thermal", "Export zynect-config.json with Bearer token from browser devtools; set site coordinates."),
        ("Pool", "Pentair + YoLink in HA; map pool-map.json; set depth offset if needed."),
        ("Pond", "Tuya sensors in HA; map pond-map.json; Zynect “Pond” sensor for temperature."),
        ("HVAC / A/C", "Nest thermostats + ha_kumo_ws mini-splits in HA; no manual mapping required."),
        ("Crestron lights", "Crestron Home integration in HA; optional room assignments on Lights page."),
        ("Outside lights", "Merge outside helpers, scripts, automations; verify no homeassistant.start on mode automations."),
        ("Remote access", "Place ha-config.json and all map JSON on HA www folder; use view.html for read-only links."),
    ]
    add_table(doc, ["Feature", "Setup steps"], checklist)

    add_heading(doc, "12. Polling & Refresh Intervals", level=2)
    add_table(
        doc,
        ["Data", "Interval", "Notes"],
        [
            ["General HA state", "5 seconds", "Shades, pool, pond, outside, HVAC, weather"],
            ["Crestron lights/scenes", "2 seconds", "Calls homeassistant.update_entity before read"],
            ["AlsoEnergy PV cache", "10 minutes", "Daylight hours only; written by custom component"],
            ["Enphase PowerPack cache", "2 minutes", "24/7; written by custom component"],
            ["Zynect (pond temp, solar thermal page)", "30 seconds default", "Configurable in zynect-config.json"],
            ["Shade cache sync (optional)", "15 minutes", "sync_caches.py automation on HA"],
        ],
    )

    add_heading(doc, "13. Key Source Files", level=2)
    add_table(
        doc,
        ["Area", "Path"],
        [
            ["App routing", "src/App.tsx"],
            ["Home layout", "src/pages/HomePage.tsx"],
            ["Global state / polling", "src/data/HouseContext.tsx"],
            ["HA client", "src/ha/client.ts"],
            ["Read-only mode", "src/dashboardMode.ts, view.html"],
            ["Deploy script", "scripts/deploy-ha.ps1"],
            ["HA package", "homeassistant/packages/home-dashboard.yaml"],
            ["Snippet docs", "homeassistant/snippets/README.md"],
            ["Cache sync", "homeassistant/dashboard_sync/sync_caches.py"],
        ],
    )

    doc.add_page_break()
    add_para(
        doc,
        "Document generated from the home-dashboard repository. "
        "Re-run scripts/generate-dashboard-doc.py after significant changes to regenerate this file.",
        bold=False,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
